import os
import config
import openpyxl
import datetime
from docx import Document

# ====================================================================================================================================================================================================================================================================================

class DeveloperTestInfo:
    def __init__(self, test_content, developer, approver):
        self.test_content = test_content
        self.developer = developer
        self.approver = approver
 
# ====================================================================================================================================================================================================================================================================================

def create_directory(directory_path):
    if not os.path.exists(directory_path):
        os.makedirs(directory_path)
        print(f"'{directory_path}' 생성")
    else:
        print(f"'{directory_path}' 이미 존재")

def get_file_names() -> list[str]:
    wb = openpyxl.load_workbook(f"{config.doc_path}/{config.base_info_file_name}")
    sheet = wb.active

    # 컬럼 데이터 읽기
    a_column = sheet['A']

    # 데이터를 저장할 리스트 초기화
    file_names = []

    # 열의 길이 중 작은 길이까지만 반복
    for a_cell in a_column[2:]: # 3행부터 테스트 데이터들 존재
        # 둘 중 하나라도 값이 비어 있으면 반복 종료
        if a_cell.value is None:
            break
        
        file_names.append(a_cell.value)

    # 파일 닫기
    wb.close()
    
    return file_names

def get_values() -> dict[str, str]:
    wb = openpyxl.load_workbook(f"{config.doc_path}/{config.base_info_file_name}")
    sheet = wb.active
    
    data_map = {}

    # 3행부터 데이터 읽기
    for row in sheet.iter_rows(min_row=3, values_only=True):
        key = row[4]  # 키 (0부터 시작하므로 5열에 해당)
        data = row[5]  # 데이터

        # 데이터가 비어있으면 중단
        if key is None or data is None:
            break

        # 딕셔너리에 데이터 추가
        data_map[key] = data

    # 파일 닫기
    wb.close()

    return data_map

def get_developer_test_infos() -> list[DeveloperTestInfo]:
    wb = openpyxl.load_workbook(f"{config.doc_path}/{config.base_info_file_name}")
    sheet = wb.active
    
    content_column = sheet['I']
    developer_column = sheet['J']

    datas = []

    # 열의 길이 중 작은 길이까지만 반복
    for content_cell, developer_cell in zip(content_column[2:], developer_column[2:]): # 3행부터 읽기
        # 둘 중 하나라도 값이 비어 있으면 반복 종료
        if content_cell.value is None or developer_cell.value is None:
            break

        datas.append(DeveloperTestInfo(test_content=content_cell.value, developer=developer_cell.value, approver="권확인자"))

    return datas

def generate_excel(file_name: str, values: dict[str, str], developer_test_infos: list[DeveloperTestInfo]):
    excel_path = f"{config.doc_path}/{file_name}"
    wb = openpyxl.load_workbook(excel_path)
    
    # 모든 시트에 대해 반복
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        
        # # 시트 내의 모든 행에 대해 반복
        # for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=1, max_col=sheet.max_column):
        #     for cell in row:
        #         # 셀의 값이 딕셔너리의 키에 해당하는 경우, 값을 딕셔너리의 값으로 업데이트
        #         if cell.value in values:
        #             cell.value = values[cell.value]
        # 시트 내의 모든 행에 대해 반복
        for row_index, row in enumerate(sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=1, max_col=sheet.max_column)):
            for col_index, cell in enumerate(row):
                # 셀의 값이 딕셔너리의 키에 해당하는 경우, 값을 딕셔너리의 값으로 업데이트
                if cell.value in values:
                    cell.value = values[cell.value]
                
                # 테스트 데이터 업데이트
                if cell.value == "{testData}":
                    # 테스트 문서 데이터 가져오기
                    # developer_test_infos = get_developer_test_infos()
                    # for developer_test_info in developer_test_infos:
                    #     print(f'테스트내용 : {developer_test_info.test_content}, 개발자명 : {developer_test_info.developer}')
                    
                    for test_info_row_index, test_info in enumerate(developer_test_infos):
                        sheet.cell(row=row_index+test_info_row_index+1, column=col_index+1, value=test_info.test_content)
                        sheet.cell(row=row_index+test_info_row_index+1, column=col_index+2, value=test_info.developer)
                        # sheet.cell(row=test_info_row_index+1, column=col_index+3, value=todo)
                        # sheet.cell(row=test_info_row_index+1, column=col_index+4, value=todo)
                        
    # 파일 저장
    wb.save(f"{config.generated_doc_path}/{file_name}")
    
    # 파일 닫기
    wb.close()
    
def generate_docx(file_name: str, values: dict[str, str]):
    docx_path = f"{config.doc_path}/{file_name}"
    docx = Document(docx_path)

    # 단락 반복
    for para in docx.paragraphs:
        # 단락 내에서 딕셔너리의 키를 찾아서 값을 딕셔너리의 값으로 업데이트
        for key, value in values.items():
            # value가 date 타입이면 string 으로 변경
            if isinstance(value, datetime.date):
                value = value.strftime("%Y-%m-%d")
            para.text = para.text.replace(key, value)

    # 테이블 반복
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in values.items():
                    if isinstance(value, datetime.date):
                        value = value.strftime("%Y-%m-%d")
                    cell.text = cell.text.replace(key, value)

    # 파일 저장
    docx.save(f"{config.generated_doc_path}/{file_name}")
    
# ====================================================================================================================================================================================================================================================================================

if __name__ == "__main__":
    # 파일명 리스트 가져오기
    file_names = get_file_names()
    for file_name in file_names:
        print(f'파일명 : {file_name}')
    
    print("===================================================")

    # 변수 맵 가져오기
    values = get_values()
    for key, value in values.items():
        print(f'변수명 : {key}, 데이터 : {value}')
        
    print("===================================================")
    
    # 테스트 문서 데이터 가져오기
    developer_test_infos = get_developer_test_infos()
    for developer_test_info in developer_test_infos:
        print(f'테스트내용 : {developer_test_info.test_content}, 개발자명 : {developer_test_info.developer}')
        
    print("===================================================")
    
    # 생성 파일 저장할 폴더 생성
    create_directory(config.generated_doc_path)
    
    # 파일 생성
    for file_name in file_names:
        if file_name.endswith(".xlsx"):
            # 엑셀 파일 생성
            generate_excel(file_name, values, developer_test_infos)
        elif file_name.endswith(".docx"):
            # 워드 파일 생성
            generate_docx(file_name, values) 
            
            
            
            
            
            
            